"""
Mishnah Structural Marker Extractor — v2.1.4

Extracts structural markers from 'The Whole Structured Mishnah for pdf.docx'
into JSON entries compatible with mishnah_db.json.

The Word document encodes structural markers via named character styles (not
direct colors). Each style maps to a CSS class used on the chaver.com site:

    Horizontal10      -> horizontal1     (blue #3399FF)
    Horizontal2       -> horizontal2     (teal #008080)
    Horizontal3       -> horizontal3     (dark cyan #008B8B)
    Vertical1         -> vertical1       (brown #8B4513)
    InternalParallel  -> internalparallel (dark red #C00000)
    Closure           -> closure         (purple #77206D)
    Ciasm1            -> ciasm1          (violet #7030A0)
    Ciasm2            -> ciasm2          (violet #7030A0, underlined)

Changes in v2.1:
    - Fixed hebrew_num to use proper gematria (not letter indices)
    - Added reversed-header matching (some tables have perek first, masekhet last)
    - Fixed key-name mismatches (oktzin, tevulyom, avodazara, tahorot)
    - Added alternate Hebrew spellings (eduyot, niddah, middot, eruvin, kiddushin)

Changes in v2.1.1:
    - Cell-order fix: reversed-header tables now produce cells in א-first order
      (matching JSON convention). Standard-header tables unchanged.
    - Added alternate Hebrew spelling: קינים -> kinnim

Changes in v2.1.2 (superseded by v2.1.3):
    - Position-based TOP/MIDDLE/BOTTOM header heuristic — too aggressive.
      174 firings across the corpus, including 49 over-fires that lost real
      Mishnah content (e.g., berakhot_1's mishnah ג).
    - Graceful missing-chapter handling (kept in v2.1.3).
    - extract_all_chapters_from_json() driver (kept in v2.1.3).

Changes in v2.1.3 (kept in v2.1.4):
    - Precise header-row rule with two conditions on Subunit row-number
      signals (see _classify_special_rows).
    - Tractate aliases for bavakamma/bavametzia/bavabatra/oholot/rosh_hashanah.
    - Sequential row renumbering after header absorption.
    - Enriched duplicate-key reporting.

Changes in v2.1.4:
    - TABLE_OVERRIDES: per-chapter dict that pins a specific docx table
      index when multiple tables resolve to the same (tractate, chap) key.
      Used for chapters where last-write-wins picks the wrong one.
    - sotah_9 split: when the docx contains tables for "מסכת סוטה פרק ט חלק א"
      and "...חלק ב", they are keyed as sotah_9a and sotah_9b respectively
      (the unified sotah_9 key is intentionally not produced). The
      extract_all_chapters_from_json driver suppresses a placeholder for
      the live JSON's sotah_9 key and inherits its source_url for both
      halves, attaching a chapter_part_he field to each half.

Usage:
    python mishnah_extractor_v2.py <docx_path> [chapter_key] [output_path]
    python mishnah_extractor_v2.py <docx_path> --all all.json
    python mishnah_extractor_v2.py <docx_path> --from-json <live.json> <out.json>

Examples:
    python mishnah_extractor_v2.py "Mishnah.docx" megillah_1 out.json
    python mishnah_extractor_v2.py "Mishnah.docx" --all all.json
    python mishnah_extractor_v2.py "Mishnah.docx" --from-json mishnah_db.json staged.json
"""

__version__ = "2.1.4"


# v2.1.4 — Per-chapter table overrides for duplicate-key cases where the
# default last-write-wins behavior picks the wrong docx table. Keys are
# "{eng_tractate}_{chap_num}" strings; values are "ti=N" identifying the
# 0-indexed docx table to use. If the docx is re-edited and table indices
# shift, the overrides below need to be reviewed.
TABLE_OVERRIDES = {
    "shabbat_22": "ti=110",  # earlier of two tables; 3×2 with 10 markers
}


# v2.1.4 — sotah_9 split keys. When a docx table matches sotah ch.9, we
# look for "חלק א" / "חלק ב" in its chapter_text and key it accordingly,
# bypassing the unified "sotah_9" key.
_SOTAH9_PART_A = "חלק א"
_SOTAH9_PART_B = "חלק ב"

import json, re, sys, os
from typing import Dict, List, Optional, Any

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx required. pip install python-docx")
    sys.exit(1)


STYLE_TO_MARKER: Dict[str, str] = {
    "Horizontal10": "horizontal1",
    "Horizontal2": "horizontal2",
    "Horizontal3": "horizontal3",
    "Vertical1": "vertical1",
    "InternalParallel": "internalparallel",
    "Closure": "closure",
    "Ciasm1": "ciasm1",
    "Ciasm2": "ciasm2",
}

SUBUNIT_STYLE = "Subunit"


def get_run_style(run_el) -> Optional[str]:
    """Get character style name from a run XML element."""
    rpr = run_el.find(qn('w:rPr'))
    if rpr is None:
        return None
    rstyle = rpr.find(qn('w:rStyle'))
    if rstyle is None:
        return None
    return rstyle.get(qn('w:val'))


def extract_run_text(run_el) -> str:
    """Extract text including <w:br/>, <w:tab/>, <w:cr/> from a run."""
    parts = []
    for child in run_el:
        tag = child.tag.split('}')[-1]
        if tag == 't':
            parts.append(child.text or '')
        elif tag == 'br':
            parts.append('\n')
        elif tag == 'tab':
            parts.append('\t')
        elif tag == 'cr':
            parts.append('\n')
    return ''.join(parts)


def extract_chapter(table) -> Dict:
    """Extract a chapter's structure and markers from a Word table.

    v2.1.1: Detects reversed-header tables (פרק in first cell) and reverses
    cell order so that column א is always at index 0 in the output.

    v2.1.3: After row extraction, applies the precise header rule
    (see _classify_special_rows): a 1-cell row becomes a 'header' on the
    next data row only when (a) its label is a standalone row-number >= 2
    AND (b) the next data row's cell labels all share that row-number
    prefix. Wide rows that fail either condition are preserved as data
    rows. After header absorption, remaining rows are renumbered 1..N.
    The header-rule firing log is returned under '_classifier_log'
    (transient field; callers strip before persistence).
    """
    # Detect header orientation early
    header_cells = table.rows[0].cells
    c0 = header_cells[0].text.strip()
    cl = header_cells[-1].text.strip()
    is_reversed_header = 'פרק' in c0 and 'מסכת' not in c0

    # Assign tractate/chapter from the correct header cells
    if is_reversed_header:
        tractate_he = cl
        chapter_text = c0
    else:
        tractate_he = c0
        chapter_text = cl
    chapter_he_match = re.search(r'פרק\s+(.+)', chapter_text)
    chapter_he = chapter_he_match.group(1).strip() if chapter_he_match else ""

    rows_data = []
    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:
            continue
        cells_data = []
        grid_col = 1
        tr_el = row._tr
        seen_tcs = set()
        for tc in tr_el.findall(qn('w:tc')):
            tc_id = id(tc)
            if tc_id in seen_tcs:
                continue
            seen_tcs.add(tc_id)
            tc_pr = tc.find(qn('w:tcPr'))
            colspan = 1
            if tc_pr is not None:
                gs = tc_pr.find(qn('w:gridSpan'))
                if gs is not None:
                    colspan = int(gs.get(qn('w:val'), '1'))
            cell_runs = []
            for para in tc.findall(qn('w:p')):
                for run_el in para.findall(qn('w:r')):
                    style = get_run_style(run_el)
                    text = extract_run_text(run_el)
                    if not text:
                        continue
                    marker = STYLE_TO_MARKER.get(style) if style else None
                    cell_runs.append({"text": text, "marker": marker, "_style": style})
            if not cell_runs:
                grid_col += colspan
                continue
            cell_dict = _parse_cell(cell_runs, row_idx, grid_col, colspan)
            if cell_dict:
                cells_data.append(cell_dict)
            grid_col += colspan

        # v2.1.1: Reverse cell order for reversed-header tables so column א
        # is always at array index 0 (matching the JSON convention).
        if is_reversed_header and cells_data:
            cells_data.reverse()
            # Reassign position col indices after reversal
            col = 1
            for cell in cells_data:
                cell['position']['col'] = col
                col += cell['position']['colspan']

        if cells_data:
            rows_data.append({"row_num": row_idx, "cells": cells_data})

    # v2.1.3: apply the precise header rule (replaces v2.1.2's position heuristic)
    rows_data, classifier_log = _classify_special_rows(rows_data)

    # v2.1.3: renumber data rows sequentially starting at 1 (matches live JSON)
    for new_idx, r in enumerate(rows_data, start=1):
        r['row_num'] = new_idx
        for cell in r.get('cells', []):
            pos = cell.get('position')
            if isinstance(pos, dict):
                pos['row'] = new_idx

    shape = [[c['position']['colspan'] for c in r['cells']] for r in rows_data]
    result = {"tractate_he": tractate_he, "chapter_he": chapter_he, "shape": shape, "rows": rows_data}
    if classifier_log:
        result["_classifier_log"] = classifier_log
    return result


# Hebrew-numeral helpers used by the v2.1.3 header rule.

_HEB_LETTER_TO_INT = {
    'א': 1, 'ב': 2, 'ג': 3, 'ד': 4, 'ה': 5,
    'ו': 6, 'ז': 7, 'ח': 8, 'ט': 9, 'י': 10,
    'כ': 20, 'ל': 30, 'מ': 40, 'נ': 50,
}


def _hebrew_letters_to_int(s):
    """Convert a Hebrew-letter numeral like 'ב', 'יב', 'טו', 'טז' to int.

    Returns None if the string is not a recognized Hebrew numeral.
    Supports values 1..49 (covers the largest Mishnah chapter row counts).
    """
    if not s:
        return None
    # Special religious-convention forms
    if s == 'טו':
        return 15
    if s == 'טז':
        return 16
    # Single letter
    if len(s) == 1:
        return _HEB_LETTER_TO_INT.get(s)
    # Two letters: tens + ones
    if len(s) == 2:
        tens = _HEB_LETTER_TO_INT.get(s[0])
        ones = _HEB_LETTER_TO_INT.get(s[1])
        if tens is None or ones is None:
            return None
        if tens not in (10, 20, 30, 40) or ones >= 10:
            return None
        return tens + ones
    return None


def _parse_subunit_signal(label):
    """v2.1.3 — Is `label` a standalone row-number signal for the header rule?

    Returns the integer row number (>= 2) if `label` is:
      - a plain Arabic-digit string like "2", "3", "12" (value >= 2), OR
      - a Hebrew-letter numeral like "ב", "ג", "טו" (value >= 2).
    Returns None otherwise.

    Specifically excludes:
      - empty/whitespace
      - "1" or "א" (row-1 wide cells are content, not headers)
      - cell labels with a column suffix (e.g., "2א", "12ב")
      - any string containing whitespace or punctuation
    """
    if not label:
        return None
    s = label.strip()
    if not s:
        return None
    if s.isdigit():
        n = int(s)
        return n if n >= 2 else None
    # Hebrew-letter form
    if all('א' <= ch <= 'ת' for ch in s):
        n = _hebrew_letters_to_int(s)
        if n is not None and n >= 2:
            return n
    return None


def _parse_cell_label_row(label):
    """Extract the row-number prefix from a cell label like '2א', '12ב', '6'.

    Returns int or None.
    """
    if not label:
        return None
    s = label.strip()
    m = re.match(r'^(\d+)', s)
    if m:
        return int(m.group(1))
    # Hebrew-letter prefix: take leading Hebrew-letter run that resolves to a
    # known numeral.
    for cut in (2, 1):
        if len(s) >= cut:
            n = _hebrew_letters_to_int(s[:cut])
            if n is not None:
                return n
    return None


def _classify_special_rows(rows_data):
    """v2.1.3 — Precise header-row detection rule.

    A 1-cell row R becomes a 'header' on the next data row D iff:
      (1) R's single cell has a label that is a standalone row-number
          signal (digit or Hebrew letter representing N >= 2 — see
          _parse_subunit_signal), AND
      (2) every cell in D has a label whose row-number prefix equals N.

    When both conditions hold, R's cell `text` is attached as a string on
    D under the key 'header', and R is removed from rows_data. Otherwise R
    is preserved as a regular data row.

    Args:
        rows_data: list of {'row_num': int, 'cells': list[cell-dict]}.

    Returns:
        (new_rows, log)
            new_rows: rows_data with header rows absorbed.
            log: list of {'row_num', 'dest_row_num', 'signal',
                          'text_preview'} for each firing.
    """
    if not rows_data:
        return rows_data, []

    log = []
    to_drop = set()
    n = len(rows_data)

    for i, row in enumerate(rows_data):
        cells = row.get('cells', [])
        if len(cells) != 1:
            continue  # only wide 1-cell rows are candidates

        wide_cell = cells[0]
        signal = _parse_subunit_signal(wide_cell.get('label', ''))
        if signal is None:
            continue  # condition (1) failed

        # Find the next row whose cells are all real data cells (skip drops
        # earlier in the loop — none yet at this point since we go in order).
        target_idx = None
        for j in range(i + 1, n):
            if j in to_drop:
                continue
            if len(rows_data[j].get('cells', [])) >= 1:
                target_idx = j
                break
        if target_idx is None:
            continue  # nothing to attach to

        target_row = rows_data[target_idx]
        target_cells = target_row.get('cells', [])
        # The next row must have at least 2 cells (it must be a real data row,
        # not another wide row).
        if len(target_cells) < 2:
            continue

        # Condition (2): all target cells' label row-num prefix equals signal.
        all_match = True
        for c in target_cells:
            c_row_num = _parse_cell_label_row(c.get('label', ''))
            if c_row_num != signal:
                all_match = False
                break
        if not all_match:
            continue

        # Fire: attach text as header, drop the wide row.
        target_row['header'] = wide_cell.get('text', '')
        to_drop.add(i)
        log.append({
            'row_num': row.get('row_num'),
            'dest_row_num': target_row.get('row_num'),
            'signal': signal,
            'text_preview': (wide_cell.get('text', '') or '').replace('\n', ' ').strip()[:80],
        })

    new_rows = [r for idx, r in enumerate(rows_data) if idx not in to_drop]
    return new_rows, log


def _parse_cell(runs, row_num, grid_col, colspan):
    """Parse a cell's runs into a structured cell dict."""
    if not runs:
        return None
    label = ""
    if runs[0].get('_style') == SUBUNIT_STYLE:
        label = runs[0]['text'].strip()
    elif runs:
        first = runs[0]['text'].strip()
        if re.match(r'^\d+[א-ת]$', first):
            label = first
    output_runs = [{"text": r["text"], "marker": r["marker"]} for r in runs]
    markers = [{"type": r["marker"], "text": r["text"]} for r in runs if r["marker"]]
    full_text = ''.join(r['text'] for r in runs)
    subdivisions = _detect_subdivisions(runs)
    cell_dict = {
        "label": label,
        "position": {"row": row_num, "col": grid_col, "colspan": colspan},
        "text": full_text,
        "runs": output_runs,
        "markers": markers,
    }
    if subdivisions:
        cell_dict["subdivisions"] = subdivisions
    return cell_dict


def _detect_subdivisions(runs):
    """Detect A/B/C subdivisions within a cell."""
    sub_indices = []
    for i, r in enumerate(runs):
        if (r.get('_style') == SUBUNIT_STYLE and
            re.match(r'^[A-Z]$', r['text'].strip())):
            sub_indices.append(i)
    if not sub_indices:
        return []
    subdivisions = []
    for idx, sub_idx in enumerate(sub_indices):
        sub_label = runs[sub_idx]['text'].strip()
        end_idx = sub_indices[idx + 1] if idx + 1 < len(sub_indices) else len(runs)
        start = sub_idx + 1
        if start < end_idx and runs[start]['text'] == ' ' and runs[start]['marker'] is None:
            start += 1
        sub_text = ''.join(r['text'] for r in runs[start:end_idx])
        sub_markers = [{"type": r["marker"], "text": r["text"]}
                       for r in runs[start:end_idx] if r["marker"]]
        mishnah_num = None
        num_match = re.match(r'\(?([א-ת]+)\)?', sub_text.lstrip())
        if num_match:
            mishnah_num = num_match.group(1)
        subdivisions.append({
            "label": sub_label, "text": sub_text,
            "markers": sub_markers, "mishnah_num_he": mishnah_num
        })
    return subdivisions


def _hebrew_chapter_num_to_str(num):
    """Convert chapter number to Hebrew numeral (gematria).

    Uses standard Hebrew numeral conventions:
    - ones: א=1 through ט=9
    - tens: י=10, כ=20, ל=30
    - 15=טו (not יה), 16=טז (not יו) — religious convention
    """
    if num <= 0:
        return str(num)
    ones = ['', 'א', 'ב', 'ג', 'ד', 'ה', 'ו', 'ז', 'ח', 'ט']
    tens = ['', 'י', 'כ', 'ל']
    if num == 15:
        return 'טו'
    if num == 16:
        return 'טז'
    result = ''
    if num >= 10:
        result += tens[num // 10]
        num = num % 10
    result += ones[num]
    return result


TRACTATE_NAMES = {
    "מגילה": "megillah",
    "מגלה": "megillah",
    "ברכות": "berakhot",
    "שבת": "shabbat",
    "עירובין": "eruvin",
    "ערובין": "eruvin",
    "פסחים": "pesachim",
    "שקלים": "shekalim",
    "יומא": "yoma",
    "סוכה": "sukkah",
    "ביצה": "beitzah",
    "ראש השנה": "rosh_hashanah",
    "תענית": "taanit",
    "מועד קטן": "moed_katan",
    "חגיגה": "chagigah",
    "יבמות": "yevamot",
    "כתובות": "ketubot",
    "נדרים": "nedarim",
    "נזיר": "nazir",
    "סוטה": "sotah",
    "גיטין": "gittin",
    "קידושין": "kiddushin",
    "קדושין": "kiddushin",
    "בבא קמא": "bavakamma",
    "בבא מציעא": "bavametzia",
    "בבא בתרא": "bavabatra",
    "סנהדרין": "sanhedrin",
    "מכות": "makkot",
    "שבועות": "shevuot",
    "עדויות": "eduyot",
    "עדיות": "eduyot",
    "עבודה זרה": "avodazara",
    "אבות": "avot",
    "הוריות": "horayot",
    "זבחים": "zevachim",
    "מנחות": "menachot",
    "חולין": "chullin",
    "בכורות": "bekhorot",
    "ערכין": "arakhin",
    "תמורה": "temurah",
    "כריתות": "keritot",
    "מעילה": "meilah",
    "תמיד": "tamid",
    "מדות": "middot",
    "מידות": "middot",
    "קנים": "kinnim",
    "קינים": "kinnim",
    "כלים": "kelim",
    "אהלות": "oholot",
    "אהילות": "oholot",
    "נגעים": "negaim",
    "פרה": "parah",
    "טהרות": "tahorot",
    "מקואות": "mikvaot",
    "נדה": "niddah",
    "נידה": "niddah",
    "מכשירין": "makhshirin",
    "זבים": "zavim",
    "טבול יום": "tevulyom",
    "ידים": "yadayim",
    "עוקצין": "oktzin",
    "עוקצים": "oktzin",
    "פאה": "peah",
    "דמאי": "demai",
    "כלאים": "kilayim",
    "שביעית": "sheviit",
    "תרומות": "terumot",
    "מעשרות": "maasrot",
    "מעשר שני": "maaser_sheni",
    "חלה": "challah",
    "ערלה": "orlah",
    "בכורים": "bikkurim",
}


def _match_tractate(hebrew_text):
    """Match Hebrew tractate text to English key."""
    for heb, eng in TRACTATE_NAMES.items():
        if heb in hebrew_text:
            return eng
    return None


def build_table_index(doc):
    """v2.1.2 — Pre-index docx tables by (english_tractate_name, chapter_num).

    Handles both standard ("מסכת X" first, "פרק Y" last) and reversed-header
    ("פרק Y" first, "מסכת X" last) table layouts. Tables without a
    recognizable header are skipped.

    Returns:
        (index, duplicates)
            index: dict {(tractate_eng, chap_num): (table, table_idx)}.
                   Later occurrences of the same key overwrite earlier ones.
            duplicates: list of (key, table_idx) for any duplicate keys
                        encountered (for the anomaly report).
    """
    index = {}
    duplicates = []
    for ti, table in enumerate(doc.tables):
        if len(table.rows) < 2:
            continue
        hcells = table.rows[0].cells
        if len(hcells) < 2:
            continue
        c0 = (hcells[0].text or '').strip()
        cl = (hcells[-1].text or '').strip()

        # Decide which header cell holds the tractate and which holds the chapter.
        # Strict standard / reversed first; then a tolerant fallback that handles
        # docx typos (e.g., nazir_8's c0 is "סכת נזיר" — missing the leading מ).
        is_reversed = 'פרק' in c0 and 'מסכת' not in c0
        tractate_text = chapter_text = None
        if is_reversed and 'מסכת' in cl:
            tractate_text, chapter_text = cl, c0
        elif 'מסכת' in c0 and 'פרק' in cl:
            tractate_text, chapter_text = c0, cl
        elif 'פרק' in cl and _match_tractate(c0):
            # Tolerant: c0 lacks "מסכת" prefix but contains a known tractate name.
            tractate_text, chapter_text = c0, cl
        elif 'פרק' in c0 and _match_tractate(cl):
            tractate_text, chapter_text = cl, c0
        else:
            continue

        eng = _match_tractate(tractate_text)
        if not eng:
            continue
        chap_match = re.search(r'פרק\s+([א-ת]+)', chapter_text)
        if not chap_match:
            continue
        chap_he = chap_match.group(1).strip()
        chap_num = None
        for n in range(1, 50):
            if _hebrew_chapter_num_to_str(n) == chap_he:
                chap_num = n
                break
        if chap_num is None:
            continue

        # Compute the string key. v2.1.4 splits sotah_9 into sotah_9a / sotah_9b
        # based on "חלק א" / "חלק ב" in chapter_text; the unified sotah_9 key
        # is intentionally not produced.
        if eng == "sotah" and chap_num == 9:
            if _SOTAH9_PART_A in chapter_text:
                key_str = "sotah_9a"
            elif _SOTAH9_PART_B in chapter_text:
                key_str = "sotah_9b"
            else:
                key_str = f"{eng}_{chap_num}"
        else:
            key_str = f"{eng}_{chap_num}"

        if key_str in index:
            prior_ti = index[key_str][1]
            prior_meta = _table_summary(doc.tables[prior_ti])
            new_meta = _table_summary(table)
            duplicates.append({
                'key': key_str,
                'prior_table_idx': prior_ti,
                'prior_shape': prior_meta['shape'],
                'prior_cell_count': prior_meta['cell_count'],
                'prior_marker_count': prior_meta['marker_count'],
                'new_table_idx': ti,
                'new_shape': new_meta['shape'],
                'new_cell_count': new_meta['cell_count'],
                'new_marker_count': new_meta['marker_count'],
            })
            # v2.1.4 — honor TABLE_OVERRIDES on duplicate-key clashes.
            override_target = TABLE_OVERRIDES.get(key_str)
            if override_target is not None:
                m = re.match(r'ti=(\d+)$', override_target)
                if m:
                    target_ti = int(m.group(1))
                    if ti == target_ti:
                        index[key_str] = (table, ti)
                    elif prior_ti == target_ti:
                        pass  # keep prior; do not overwrite
                    else:
                        # Neither candidate matches the override; fall back
                        # to last-write-wins for safety.
                        index[key_str] = (table, ti)
                    continue
        index[key_str] = (table, ti)
    return index, duplicates


def _table_summary(table):
    """Compute (shape, cell_count, marker_count) for a single table.

    Used for duplicate reporting; runs extract_chapter once to get counts.
    """
    chap = extract_chapter(table)
    chap.pop('_classifier_log', None)
    cell_count = sum(len(r.get('cells', [])) for r in chap.get('rows', []))
    marker_count = 0
    for r in chap.get('rows', []):
        for c in r.get('cells', []):
            marker_count += len(c.get('markers', []) or [])
            for s in c.get('subdivisions', []) or []:
                marker_count += len(s.get('markers', []) or [])
    return {
        'shape': chap.get('shape'),
        'cell_count': cell_count,
        'marker_count': marker_count,
    }


def extract_all_chapters_from_json(docx_path, live_json_path):
    """v2.1.4 — Iterate live JSON keys, extract or emit placeholder.

    For each live key (excluding _meta), match against the v2.1.4 index
    (which uses string keys). Match by exact key string; v2.1.4 produces
    keys like 'megillah_1', and the split sotah keys 'sotah_9a'/'sotah_9b'.

    Special handling:
      - The live JSON's 'sotah_9' key is intentionally suppressed: no
        placeholder is emitted. Fresh 'sotah_9a' / 'sotah_9b' entries are
        appended after the live-key pass, inheriting source_url and
        seder metadata from the live sotah_9 entry. Each gets a new
        chapter_part_he field ("חלק א" / "חלק ב") to distinguish them.

    Returns:
        (results, classifier_logs, missing_keys, duplicates, table_idx_used)
    """
    doc = Document(docx_path)
    with open(live_json_path, 'r', encoding='utf-8') as f:
        live = json.load(f)

    index, duplicates = build_table_index(doc)

    results = {}
    classifier_logs = {}
    missing_keys = []
    table_idx_used = {}

    live_chapter_keys = sorted(k for k in live.keys() if k != '_meta')
    sotah_9_live = live.get('sotah_9')  # used to seed sotah_9a/9b metadata

    for key in live_chapter_keys:
        # v2.1.4 — suppress the live sotah_9 key; sotah_9a/9b emitted below.
        if key == 'sotah_9':
            continue

        live_entry = live[key] if isinstance(live.get(key), dict) else {}

        if key in index:
            table, ti = index[key]
            chapter_data = extract_chapter(table)
            chapter_data['tractate_en'] = live_entry.get('tractate_en')
            chapter_data['seder_he'] = live_entry.get('seder_he')
            chapter_data['seder_en'] = live_entry.get('seder_en')
            chapter_data['chapter_num'] = live_entry.get('chapter_num')
            chapter_data['source_url'] = live_entry.get('source_url')
            log = chapter_data.pop('_classifier_log', None)
            if log:
                classifier_logs[key] = log
            table_idx_used[key] = ti
            results[key] = chapter_data
        else:
            placeholder = {
                'tractate_he': live_entry.get('tractate_he'),
                'tractate_en': live_entry.get('tractate_en'),
                'seder_he': live_entry.get('seder_he'),
                'seder_en': live_entry.get('seder_en'),
                'chapter_num': live_entry.get('chapter_num'),
                'chapter_he': live_entry.get('chapter_he'),
                'source_url': live_entry.get('source_url'),
                'shape': [],
                'rows': [],
                '_missing_from_docx': True,
            }
            results[key] = placeholder
            missing_keys.append(key)

    # v2.1.4 — emit sotah_9a and sotah_9b as fresh keys (if found in docx).
    sotah_meta = {
        'tractate_he': (sotah_9_live or {}).get('tractate_he') or 'מסכת סוטה',
        'tractate_en': (sotah_9_live or {}).get('tractate_en') or 'Sotah',
        'seder_he': (sotah_9_live or {}).get('seder_he') or 'נשים',
        'seder_en': (sotah_9_live or {}).get('seder_en') or 'Nashim',
        'chapter_num': 9,
        'chapter_he': 'ט',
        'source_url': (sotah_9_live or {}).get('source_url'),
    }
    for part_key, part_he in (('sotah_9a', _SOTAH9_PART_A),
                              ('sotah_9b', _SOTAH9_PART_B)):
        if part_key in index:
            table, ti = index[part_key]
            chapter_data = extract_chapter(table)
            # Apply metadata in the desired order (tractate_he/en first,
            # then seder, then chapter info, then content).
            ordered = {
                'tractate_he': sotah_meta['tractate_he'],
                'tractate_en': sotah_meta['tractate_en'],
                'seder_he': sotah_meta['seder_he'],
                'seder_en': sotah_meta['seder_en'],
                'chapter_num': sotah_meta['chapter_num'],
                'chapter_he': sotah_meta['chapter_he'],
                'chapter_part_he': part_he,
                'source_url': sotah_meta['source_url'],
                'shape': chapter_data.get('shape', []),
                'rows': chapter_data.get('rows', []),
            }
            log = chapter_data.pop('_classifier_log', None)
            if log:
                classifier_logs[part_key] = log
            table_idx_used[part_key] = ti
            results[part_key] = ordered

    return results, classifier_logs, missing_keys, duplicates, table_idx_used


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    docx_path = sys.argv[1]
    if not os.path.exists(docx_path):
        print(f"ERROR: File not found: {docx_path}")
        sys.exit(1)
    print(f"Loading: {docx_path}")
    doc = Document(docx_path)
    print(f"  Tables: {len(doc.tables)}")

    if len(sys.argv) >= 3 and sys.argv[2] == '--all':
        output_path = sys.argv[3] if len(sys.argv) > 3 else 'all_chapters.json'
        print('Extracting all chapters...')
        results = {}
        for table in doc.tables:
            if len(table.rows) < 2:
                continue
            hcells = table.rows[0].cells
            if len(hcells) < 2:
                continue
            c0 = hcells[0].text.strip()
            cl = hcells[-1].text.strip()
            # Try both header orientations
            tractate_text = chapter_text = None
            if 'מסכת' in c0 and 'פרק' in cl:
                tractate_text, chapter_text = c0, cl
            elif 'פרק' in c0 and 'מסכת' in cl:
                tractate_text, chapter_text = cl, c0
            if tractate_text and chapter_text:
                chapter_data = extract_chapter(table)
                key = f"{tractate_text}|{chapter_text}"
                results[key] = chapter_data
        print(f"  Found {len(results)} chapters")
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(results, f, ensure_ascii=False, indent=2)
        print(f"  Written to: {output_path}")

    elif len(sys.argv) >= 4 and sys.argv[2] == '--from-json':
        live_json_path = sys.argv[3]
        output_path = sys.argv[4] if len(sys.argv) > 4 else 'mishnah_db_reextracted.json'
        print(f'Driving extraction from live JSON: {live_json_path}')
        if not os.path.exists(live_json_path):
            print(f"ERROR: Live JSON not found: {live_json_path}")
            sys.exit(1)
        results, logs, missing, duplicates, table_idx_used = (
            extract_all_chapters_from_json(docx_path, live_json_path))
        # Compute summary
        matched = sum(1 for v in results.values() if not v.get('_missing_from_docx'))
        n_missing = sum(1 for v in results.values() if v.get('_missing_from_docx'))
        with_markers = 0
        marker_counts = {}
        with_subdivisions = 0
        with_header = 0
        for v in results.values():
            if v.get('_missing_from_docx'):
                continue
            chap_has_markers = False
            for r in v.get('rows', []):
                if 'header' in r:
                    with_header += 1
                for c in r.get('cells', []):
                    for m in c.get('markers', []):
                        marker_counts[m['type']] = marker_counts.get(m['type'], 0) + 1
                        chap_has_markers = True
                    if c.get('subdivisions'):
                        with_subdivisions += 1
                    for s in c.get('subdivisions', []) or []:
                        for m in s.get('markers', []):
                            marker_counts[m['type']] = marker_counts.get(m['type'], 0) + 1
                            chap_has_markers = True
            if chap_has_markers:
                with_markers += 1
        output_payload = {
            '_meta': {
                'version': 'stage-a-staged',
                'extractor_version': __version__,
                'source_docx_basename': os.path.basename(docx_path),
                'live_json_basename': os.path.basename(live_json_path),
                'total_chapters': len(results),
                'matched_chapters': matched,
                'missing_chapters': n_missing,
                'chapters_with_markers': with_markers,
                'chapters_with_subdivisions': with_subdivisions,
                'chapters_with_header_row': with_header,
                'total_markers': sum(marker_counts.values()),
                'marker_counts': marker_counts,
                'classifier_logs': logs,
                'classifier_fired_count': len(logs),
                'duplicates_in_docx': duplicates,
            }
        }
        output_payload.update(results)
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(output_payload, f, ensure_ascii=False, indent=2)
        print(f"  Written: {len(results)} chapters")
        print(f"  Matched: {matched}, Missing: {n_missing}")
        print(f"  Chapters with markers: {with_markers}")
        print(f"  Total markers: {sum(marker_counts.values())}")
        print(f"  Classifier fired in: {len(logs)} chapters")
        print(f"  Output: {output_path}")

    elif len(sys.argv) >= 3:
        chapter_key = sys.argv[2]
        output_path = sys.argv[3] if len(sys.argv) > 3 else f'{chapter_key}_extracted.json'
        parts = chapter_key.rsplit('_', 1)
        if len(parts) != 2:
            print(f"ERROR: Invalid key: {chapter_key}")
            sys.exit(1)
        tractate_name = parts[0]
        chapter_num = int(parts[1])
        chapter_he = _hebrew_chapter_num_to_str(chapter_num)
        target_perek = 'פרק ' + chapter_he
        print(f"  Looking for: {chapter_key} (ch. {chapter_he})")
        found_table = None
        for table in doc.tables:
            if len(table.rows) < 2:
                continue
            hcells = table.rows[0].cells
            if len(hcells) < 2:
                continue
            c0 = hcells[0].text.strip()
            cl = hcells[-1].text.strip()
            # Try both header orientations (standard and reversed)
            tractate_text = None
            if 'מסכת' in c0 and target_perek in cl:
                tractate_text = c0
            elif target_perek in c0 and 'מסכת' in cl:
                tractate_text = cl
            if tractate_text:
                eng = _match_tractate(tractate_text)
                if eng and eng == tractate_name:
                    found_table = table
                    print(f"  Found: {c0} / {cl}")
                    break
        if found_table is None:
            print(f"ERROR: Table not found for {chapter_key}")
            sys.exit(1)
        chapter_data = extract_chapter(found_table)
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(chapter_data, f, ensure_ascii=False, indent=2)
        total_m = sum(len(c['markers']) for r in chapter_data['rows'] for c in r['cells'])
        total_c = sum(len(r['cells']) for r in chapter_data['rows'])
        sub_c = sum(1 for r in chapter_data['rows'] for c in r['cells'] if c.get('subdivisions'))
        print(f"  Cells: {total_c}, Markers: {total_m}, Subdivided: {sub_c}")
        print(f"  Written to: {output_path}")
    else:
        print(__doc__)
        sys.exit(1)


if __name__ == '__main__':
    main()
